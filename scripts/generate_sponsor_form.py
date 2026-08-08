#!/usr/bin/env python3

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_SUMMARY = [
    "An Android news-feed application implemented with Kotlin, Jetpack Compose, MVVM, Retrofit, and Room.",
    "Multi-channel feeds, explainable recommendation reasons, pull-to-refresh, pagination, offline cache fallback, article detail, and basic video playback.",
    "A Go and PostgreSQL REST backend with deterministic data, recommendation ranking, validation, health checks, and media delivery.",
    "Docker Compose deployment, automated Android and Go tests, integration and load testing, CodeQL and Semgrep SAST, Trivy scanning, and OWASP ZAP DAST.",
    "Technical documentation covering requirements, architecture, design, testing, DevSecOps, security, project management, and added value.",
    "Internal company deployment as reported by the student. Production URLs, infrastructure identifiers, internal screenshots, traffic, customer information, and user data are excluded.",
]

RATING_AREAS = [
    "Completion of the agreed project scope",
    "Technical quality and engineering approach",
    "Testing, security, and delivery discipline",
    "Professionalism and communication",
    "Business or user value of the delivered work",
    "Overall project performance",
]


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def border_cell(cell, color="808080", size="8"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(4)
    return heading


def add_grid_table(document, rows, widths=None):
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            set_cell_text(cells[column_index], value, bold=row_index == 0)
            border_cell(cells[column_index])
            if widths:
                cells[column_index].width = Inches(widths[column_index])
            if row_index == 0:
                shade_cell(cells[column_index], "D9E2E8")
            elif row_index % 2 == 0:
                shade_cell(cells[column_index], "F7F9FA")
    return table


def build_document(args):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)

    title = document.add_heading("Sponsor / Mentor Project Rating and Sign-off", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    notice = document.add_table(rows=1, cols=1)
    notice.style = "Table Grid"
    shade_cell(notice.cell(0, 0), "F3F1E9")
    set_cell_text(
        notice.cell(0, 0),
        "Student-prepared draft for Sponsor review. This is not an official NUS-ISS form. "
        "The Sponsor should correct any inaccurate statement before rating and signing. "
        "It does not replace feedback that NUS-ISS may request directly from the registered Sponsor.",
    )

    add_heading(document, "1. Pre-filled Project Information", 1)
    project_rows = [
        ["Field", "Details"],
        ["Project title", "Modern Recommendation Feed Demo"],
        ["Student", "Xu Ziyi"],
        ["Student ID", args.student_id],
        ["Internship company", args.company],
        ["Project period", "23 March 2026 to 14 August 2026"],
        ["Sponsor / Mentor name", args.mentor_name],
        ["Sponsor / Mentor job title", args.mentor_title],
        ["Company email", args.mentor_email],
    ]
    add_grid_table(document, project_rows, widths=[2.15, 5.55])

    add_heading(document, "2. Student-prepared Delivery Summary", 1)
    paragraph = document.add_paragraph("The student reports completion of the following agreed project work for Sponsor review:")
    paragraph.paragraph_format.space_after = Pt(2)
    for item in PROJECT_SUMMARY:
        bullet = document.add_paragraph(item, style="List Bullet")
        bullet.paragraph_format.left_indent = Inches(0.2)
        bullet.paragraph_format.first_line_indent = Inches(-0.12)
        bullet.paragraph_format.space_after = Pt(1)

    add_heading(document, "3. Sponsor / Mentor Rating", 1)
    document.add_paragraph(
        "Enter one number in each row. 1 = Unsatisfactory, 2 = Needs Improvement, "
        "3 = Satisfactory, 4 = Good, 5 = Excellent."
    )
    rating_rows = [["Assessment area", "Score (1-5)"]]
    rating_rows.extend([[area, ""] for area in RATING_AREAS])
    rating_table = add_grid_table(document, rating_rows, widths=[6.35, 1.35])
    for row in rating_table.rows[1:]:
        row.height = Inches(0.28)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "4. Acceptance and Deployment Confirmation", 1)
    document.add_paragraph("Enter one option number for each item.")
    confirmation_rows = [
        ["Confirmation item", "Numbered options", "Response"],
        ["Project acceptance", "1 = Accepted\n2 = Accepted with minor follow-up items\n3 = Not yet accepted", ""],
        ["Deployment status", "1 = Deployed internally\n2 = Ready for internal deployment or trial\n3 = Demonstrated but not deployed\n4 = Not confirmed", ""],
        ["Confidentiality restriction applies", "1 = Yes\n2 = No", ""],
    ]
    confirmation_table = add_grid_table(document, confirmation_rows, widths=[2.05, 4.55, 1.1])
    for row in confirmation_table.rows[1:]:
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "5. Sponsor / Mentor Sign-off", 1)
    document.add_paragraph(
        "I have reviewed the pre-filled project summary and corrected it where necessary. "
        "I confirm that the selected ratings, acceptance status, and deployment status reflect "
        "my assessment of the student's internship project."
    )
    signoff_rows = [
        ["Field", "Details"],
        ["Sponsor / Mentor name", args.mentor_name],
        ["Signature", ""],
        ["Date", ""],
    ]
    signoff_table = add_grid_table(document, signoff_rows, widths=[2.15, 5.55])
    signoff_table.rows[2].height = Inches(0.45)

    final_note = document.add_paragraph(
        "The signed form may be shown as sanitized supporting evidence in the student's NUS-ISS presentation. "
        "NUS-ISS may independently contact the registered Sponsor for formal project and individual feedback."
    )
    final_note.paragraph_format.space_before = Pt(5)
    for run in final_note.runs:
        run.italic = True
        run.font.size = Pt(8)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    parser.add_argument("--student-id", default="______________________________")
    parser.add_argument("--company", default="______________________________")
    parser.add_argument("--mentor-name", default="______________________________")
    parser.add_argument("--mentor-title", default="______________________________")
    parser.add_argument("--mentor-email", default="______________________________")
    return parser.parse_args()


if __name__ == "__main__":
    build_document(parse_args())
